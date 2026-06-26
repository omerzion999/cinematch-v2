"""
Build the Hebrew project report (Word .docx) mapped 1:1 to the course rubric's
10 sections. Reads the analysis CSVs in docs/analysis/ so the embedded tables
stay in sync with the data.

Dev-only tool (needs python-docx, not a runtime dependency).
Run from repo root:  python docs/report/build_report.py
Output: docs/report/CineMatch_Report.docx
"""

import os

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(__file__)
ANALYSIS = os.path.join(HERE, "..", "analysis")
OUT = os.path.join(HERE, "CineMatch_Report.docx")

FONT = "David"
DARK = RGBColor(0x1F, 0x38, 0x64)
LINK_HEX = "1F3864"
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

_bookmark_id = [0]


# ── Clickable links + bookmarks ───────────────────────────────────────────────

def _link_run_pr(size, rtl):
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), LINK_HEX); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    if rtl:
        rPr.append(OxmlElement("w:rtl"))
    return rPr


def add_external_hyperlink(paragraph, url, text, size=11):
    """Append a clickable external hyperlink (the URL opens in a browser)."""
    r_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run.append(_link_run_pr(size, rtl=False))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def add_bookmark(paragraph, name):
    """Mark a heading paragraph as an internal anchor target."""
    _bookmark_id[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_bookmark_id[0]))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(_bookmark_id[0]))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor, text, size=11):
    """Append a clickable internal hyperlink that jumps to a bookmark (RTL text)."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run.append(_link_run_pr(size, rtl=True))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


# ── RTL helpers ───────────────────────────────────────────────────────────────

def _rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:bidi"), {}))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn("w:rtl"), {}))
    run.font.name = FONT


def heading(doc, text, level=1, bookmark=None):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    _rtl_run(run)
    _rtl_paragraph(p)
    for r in p.runs:
        r.font.color.rgb = DARK
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    _rtl_run(run)
    _rtl_paragraph(p)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        _rtl_run(run)
        _rtl_paragraph(p)


def _cell_text(cell, text, bold, rtl):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)
    if rtl:
        _rtl_run(run)
        _rtl_paragraph(p)
    else:
        # English / numeric table: keep natural left-to-right order + alignment so
        # columns and numbers read logically (the bidi flip is what made them look
        # reversed). Force a Latin font so digits/Latin glyphs render cleanly.
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return run


def table_from_df(doc, df, headers=None, rtl=True):
    """Render a DataFrame as a table. rtl=True for Hebrew tables (RTL column
    order); rtl=False for English / numeric tables, which stay left-to-right so the
    column order and the mathematical order of values read logically."""
    cols = list(df.columns)
    display_headers = headers or cols
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    if rtl:
        t._tbl.tblPr.append(t._tbl.tblPr.makeelement(qn("w:bidiVisual"), {}))
    hdr = t.rows[0].cells
    for i, h in enumerate(display_headers):
        _cell_text(hdr[i], h, bold=True, rtl=rtl)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(cols):
            _cell_text(cells[i], row[c], bold=False, rtl=rtl)
    doc.add_paragraph()


def _csv(name):
    return pd.read_csv(os.path.join(ANALYSIS, name))


# ── Document ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)

    # Cover
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    r = title.add_run("CineMatch v2")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = DARK
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    rs = sub.add_run("סוכן AI שיחתי דו-לשוני להמלצת סדרות טלוויזיה")
    rs.font.size = Pt(18)
    _rtl_run(rs)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        doc.add_paragraph()
    for line in ["מגישים: תומר בלונד, עומר ציון",
                 "סדנת חדשנות מבוססת AI ו-ML (277302)",
                 'ד"ר ליהי רייכלסון · האקדמית תל-אביב-יפו',
                 "יוני 2026"]:
        p = doc.add_paragraph()
        rr = p.add_run(line)
        rr.font.size = Pt(13)
        _rtl_run(rr)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for label, url in [("Live app:  ", "https://cinematch-ai-v2.onrender.com/"),
                       ("GitHub:  ", "https://github.com/omerzion999/cinematch-v2")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        add_external_hyperlink(p, url, url, size=11)

    doc.add_page_break()

    # TOC (manual, with clickable internal links to each section)
    heading(doc, "תוכן עניינים", level=1)
    toc = [
        ("sec1", "1. רקע, הצהרת בעיה ויעדים"),
        ("sec2", "2. חדשנות ומקוריות"),
        ("sec3", "3. עומק טכני AI (NLP/LLM)"),
        ("sec4", "4. הנתונים (Big Data): מקורות, איסוף, ניקוי והכנה"),
        ("sec56", "5-6. בחירת מודל, מימוש והשוואה"),
        ("sec7", "7. בדיקה ואימות"),
        ("sec8", "8. שיקולים אתיים"),
        ("sec9", "9. מסקנות, יישום ומגבלות"),
        ("sec10", "10. הערכת השפעה"),
    ]
    for anchor, label in toc:
        p = doc.add_paragraph()
        _rtl_paragraph(p)
        add_internal_hyperlink(p, anchor, label)
    doc.add_page_break()

    # 1
    heading(doc, "1. רקע, הצהרת בעיה ויעדים", level=1, bookmark="sec1")
    para(doc, "הדומיין: מערכות המלצה ועיבוד שפה טבעית (NLP) בעולם הסטרימינג. ריבוי הפלטפורמות "
              "והקטלוגים יצר 'עומס בחירה' - הצופה מבזבז זמן רב בחיפוש מה לראות.")
    para(doc, "למה בחרנו בנושא: רצינו לבנות סוכן שיחתי דו-לשוני (עברית כברירת מחדל, ואנגלית) "
              "שממליץ על סדרות טלוויזיה מתוך קטלוג אצור משלנו, בשיחה טבעית ואישית.")
    para(doc, "הצהרת בעיה:", bold=True)
    bullets(doc, [
        "ממשקי חיפוש מסורתיים אינם אישיים ואינם שיחתיים.",
        "מודלי שפה גדולים נוטים 'להזות' כותרים שאינם קיימים או אינם זמינים.",
        "הרלוונטיות לאתגרי NLP/LLM הנוכחיים: עיגון (grounding) של מודל שפה במקור נתונים אמין, "
        "וטיפול בשאלות מחוץ לתחום מבלי לאבד את ההקשר.",
    ])
    para(doc, "יעדים:", bold=True)
    bullets(doc, [
        "המלצות אישיות (1 עד 3 כותרים) המבוססות אך ורק על הקטלוג שלנו (catalog-first).",
        "שיחה רציפה עם זיכרון הקשר.",
        "טיפול חכם בבקשות מחוץ לנושא באמצעות 'גישור' להמלצה רלוונטית.",
        "חוויית עברית RTL מלאה ושוטפת.",
    ])

    # 2
    heading(doc, "2. חדשנות ומקוריות", level=1, bookmark="sec2")
    bullets(doc, [
        "ארכיטקטורה היברידית: סוכן LLM להבנת כוונה וניהול שיחה, יחד עם מנוע דירוג שקוף "
        "מבוסס-תוכן מעל קטלוג אצור - כך ההמלצות לעולם אינן 'הזיה' של המודל.",
        "מנגנון 'גישור' (bridge): בקשה מחוץ לנושא ('בא לי לבשל פסטה') מתורגמת לחיפוש סדרות "
        "רלוונטיות בקטלוג (סדרות בישול) במקום סירוב.",
        "ניקוד binge_fit_score מהונדס ודירוג IMDb משוקלל-הצבעות, כדי שדירוג גבוה מעט-הצבעות "
        "לא יקפוץ מעל כותר איכותי ומוכר.",
        "אונבורדינג מבוסס-זרעים (seed picker): במקום שאלות מופשטות בלבד, המשתמש בוחר סדרות "
        "מוכרות שהוא אוהב, והמנוע ההיברידי מאתר סדרות דומות (similarity על מספר זרעים) - "
        "המלצה אישית ומדויקת יותר שגם ממחישה את המנוע.",
        "דו-לשוניות אמיתית כולל הסברים שוטפים בעברית הנוצרים על ידי המודל.",
    ])

    # 3
    heading(doc, "3. עומק טכני AI (NLP/LLM)", level=1, bookmark="sec3")
    para(doc, "המושגים שנלמדו בכיתה ויישומם הישיר בפרויקט:", bold=True)
    bullets(doc, [
        "דמיון טקסט (Text similarity, שבוע 5): מדדי Jaccard (קבוצות ז'אנר/עשור/שפה) ו-Cosine "
        "(על וקטורים נומריים) - שני המדדים שולבו למנוע ההיברידי.",
        "שיבוצי מילים ומשפטים (Word Embedding ו-NLP, שבועות 9-11): וקטורים רב-לשוניים בני 384 "
        "ממדים (paraphrase-multilingual-MiniLM-L12-v2, מנורמלי L2) על תקצירי הסדרות, ו-Cosine "
        "עליהם ללכידת דמיון סמנטי. השיבוצים חושבו מחדש לאחר השלמת התקצירים החסרים (סעיף 4).",
        "אשכול (Clustering, שבוע 4): K-Means לבניית 'פרופילי טעם'; נשמר כהשוואת מודל (סעיף 5-6).",
        "סוכני LLM ו-Generative AI (LLM Agents שבוע 8, ChatGPT שבועות 6-7): Groq "
        "llama-3.1-8b-instant לחילוץ כוונה (intent parsing), ניהול תור השיחה "
        "(chat/search/refine/swap) והסברים בשפה טבעית.",
        "זיהוי חריגות (Anomaly detection): סף מכויל לזיהוי מצב 'אין התאמה טובה'.",
    ])
    para(doc, "מה עשינו על הנתונים שלנו בהקשר זה: הנדסת מאפיינים (z-scores לדירוג/הצבעות/שנה/"
              "פופולריות, ניקוד binge_fit_score, חלוקה ל-rating_bucket ול-decade), השלמת תקצירים "
              "חסרים וחישוב שיבוצים סמנטיים על התקצירים (raw text -> 384-dim vector).")

    # 4
    heading(doc, "4. הנתונים (Big Data): מקורות, איסוף, ניקוי והכנה", level=1, bookmark="sec4")
    para(doc, "איסוף: מיזגנו ארבעה מאגרי Kaggle על סדרות טלוויזיה. לאחר מיזוג, הסרת כפילויות "
              "וניקוי התקבל קטלוג אחיד של 11,013 סדרות (22 עמודות).")
    src = pd.DataFrame({
        "מקור": ["TMDb", "IMDb top-5000", "Disney+", "IMDb (נוסף)", 'סה"כ לאחר מיזוג וניקוי'],
        "רשומות": ["8,309", "2,559", "138", "7", "11,013"],
    })
    table_from_df(doc, src)
    para(doc, "ניקוי והכנה (Pipeline):", bold=True)
    bullets(doc, [
        "מיזוג ארבעת המקורות לסכמה אחידה והסרת כפילויות לפי כותר.",
        "נרמול מחרוזות הז'אנר (genre_set_str) וטיפול בערכים חסרים.",
        "ערכים חסרים: num_seasons מאוכלס ב-75% מהרשומות; השלמה למדיאן בעת בניית מאפיינים.",
        "חישוב מאפיינים מהונדסים: z-scores, decade/decade_str, rating_bucket, binge_fit_score.",
    ])
    para(doc, "השלמת תקצירים (Overview backfill):", bold=True)
    para(doc, "מניפולציית נתונים מרכזית: ל-2,855 מהרשומות (כ-26%) לא היה תקציר, וביניהן כל "
              "הסדרות המפורסמות (Breaking Bad, Game of Thrones, The Office ועוד) - כך שהשיבוצים "
              "הסמנטיים שלהן היו חסרי משמעות. השלמנו את התקצירים החסרים ממקור ה-TMDB שלנו "
              "(tvs.csv, ~85 אלף תקצירים) בהתאמה לפי שם, ומילאנו 2,302 רשומות. הכיסוי עלה "
              "מ-74% ל-95%, ואז חושבו השיבוצים מחדש - מה שחיזק ישירות את חיפוש ה'דומה ל-X'.")
    diagram = os.path.join(HERE, "data_flow.png")
    if os.path.exists(diagram):
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(diagram, width=Inches(6.2))
        cap = para(doc, "תרשים 1: צינור עיבוד הנתונים מקצה לקצה (מקורות גולמיים עד הקטלוג והשיבוצים).", size=9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "התפלגות ז'אנרים לפי עשור (% מהכותרים בעשור):", bold=True)
    table_from_df(doc, _csv("genre_share_by_decade.csv"), rtl=False)
    para(doc, "מגמות בולטות (שינוי בנקודות אחוז, 2000s ← 2020s): דרמה ופשע במגמת עלייה; "
              "קומדיה ואנימציה במגמת ירידה.")
    para(doc, "דירוג ו-binge_fit לפי עשור:", bold=True)
    table_from_df(doc, _csv("rating_by_decade.csv"), rtl=False)
    para(doc, "אנטרופיית שאנון של תמהיל הז'אנרים יציבה סביב 3.7 ביט - הקטלוג מגוון לאורך העשורים.")

    # 5-6
    heading(doc, "5-6. בחירת מודל, מימוש והשוואה", level=1, bookmark="sec56")
    para(doc, "בחנו שתי אופציות לליבת מנוע ההמלצה:", bold=True)
    bullets(doc, [
        "אופציה א': אשכול K-Means ובחירת כותרים מהאשכול הקרוב לטעם המשתמש.",
        "אופציה ב': דירוג משוקלל של כל הקטלוג מול וקטור העדפות (ז'אנר, עידן, אורך, פופולריות) "
        "בשילוב ניקוד איכות.",
    ])
    para(doc, "ממצא מרכזי: אשכולות ה-K-Means קרסו. חמישה מתוך שמונה אשכולות הם תמהיל "
              "'דרמה+קומדיה' כמעט זהה, וציון ה-silhouette נמוך (~0.17) לכל k. לכן בחירת ז'אנר "
              "כמעט שאינה משנה את האשכול הנבחר ואת ההמלצות.")
    table_from_df(doc, _csv("clustering_profiles.csv"), rtl=False)
    para(doc, "ציוני silhouette לפי k:", bold=True)
    table_from_df(doc, _csv("clustering_silhouette.csv"), rtl=False)
    para(doc, "המסקנה: בחרנו באופציה ב' (דירוג משוקלל) כמנוע החי - שקוף, מבדל בין העדפות, "
              "וניתן להסבר. ה-K-Means נשמר במערכת ובדוח ככלי השוואה (הרובריקה מעריכה אשכול).")
    para(doc, "השוואת מדדי דמיון (Spearman על 5,000 זוגות אקראיים):", bold=True)
    table_from_df(doc, _csv("similarity_correlations.csv"), rtl=False)
    para(doc, "המתאם הנמוך בין Jaccard ל-Cosine (0.0736 בתרגיל 2; כ-0.18 בשחזור כאן) מלמד "
              "שהשיטות לוכדות אותות שונים, ולכן שילוב היברידי מצדיק את עצמו עבור 'דומה ל-X'.")
    para(doc, "שימוש שני במנוע ההיברידי - בורר הזרעים (seed picker): לאחר בחירת הז'אנר, "
              "המשתמש בוחר עד 3 סדרות מוכרות שהוא אוהב. לכל זרע מורץ המנוע ההיברידי, "
              "המועמדים ממוזגים לפי ציון ההיברידי המקסימלי, הזרעים מוסרים, ומסנני העידן/"
              "פופולריות מוחלים (עם הרפיה אם נשארו אפס). כך אותו מנוע משרת גם המלצה לפי טעם "
              "וגם 'דומה ל-X'.")
    para(doc, "פסאודו-קוד (קצה-לקצה):", bold=True)
    para(doc,
         "ONBOARDING:\n"
         "  if seeds picked -> recommend_from_seeds:\n"
         "     for each seed: hybrid(Jaccard + Cosine_num + Cosine_text)\n"
         "     merge candidates by MAX hybrid_score; drop seeds\n"
         "     apply era/popularity filters (relax to unfiltered if empty)\n"
         "  else -> rank_by_preferences(answers)\n\n"
         "rank_by_preferences(answers):\n"
         "  genre = HARD FLOOR (never relaxed; empty floor -> no-match message)\n"
         "  apply era > popularity > length; drop the weakest until >=1 survives\n"
         "  score = 0.6*binge_fit + 0.4*IMDb_weighted_rating; diversity de-dup; top 1..3\n\n"
         "CHAT (each user turn):\n"
         "  if gibberish -> ask to rephrase\n"
         "  elif off-topic & bridgeable -> search by topic keywords\n"
         "       (keep only shows ABOUT the topic: keyword in title OR >=2 keywords)\n"
         "  elif names a show ('like X') -> FRESH search seeded on X (not a swap)\n"
         "  else -> chat_turn LLM decides: search | refine | swap_slot | chat\n"
         "          (blank request -> ask one guiding question)\n"
         "  search/refine: seed -> hybrid + anomaly gate; else rank_by_preferences\n"
         "  return top 1..3 cards (catalog-only); fallback offline path if LLM down", size=9)
    para(doc, "דוגמאות תוצאה על מקבץ קטן (פלט דטרמיניסטי מאומת של שכבת המנוע):", bold=True)
    bullets(doc, [
        "פשע / חדש / להיטים ← Scam 1992, Dexter: Resurrection, The Penguin.",
        "קומדיה / קלאסי ← The Office, Friends, Freaks and Geeks.",
        'מד"ב / פנינה נסתרת ← Goblin, High School DxD, Love Alarm.',
        "זרעים Breaking Bad + Better Call Saul ← Bosch, Ezel, Peaky Blinders (הזרעים מוסרים).",
    ])

    # 7
    heading(doc, "7. בדיקה ואימות", level=1, bookmark="sec7")
    para(doc, "המערכת מכוסה ב-226 בדיקות אוטומטיות (148 pytest בצד השרת, 78 Vitest בצד הלקוח), "
              "כולל בדיקות שמוכיחות שתשובות שונות מניבות המלצות שונות בבירור.")
    para(doc, "שלושה מקרי בדיקה לפחות, כולל מקרה כשל:", bold=True)
    bullets(doc, [
        "מקרה 1: חובב פשע, סדרות חדשות ולהיטים ← שלוש סדרות פשע מ-2020 ואילך בעלות דירוג גבוה.",
        "מקרה 2: בורר זרעים - Breaking Bad + Better Call Saul ← שכנים מז'אנר פשע/דרמה "
        "(Bosch, Ezel, Peaky Blinders), כשהזרעים עצמם מוסרים מהתוצאות.",
        "מקרה 3: פנינה נסתרת בז'אנר מד\"ב ← כותרים מדורגים גבוה עם מעט הצבעות.",
        "מקרה 4 (שיחה): 'תמליץ על סדרה כמו The Office' ← חיפוש חדש מבוסס-זרע "
        "(It's Always Sunny, Parks and Recreation, The IT Crowd), ולא החלפה של המלצות קודמות.",
        "מקרה 5 (גישור מחוץ לנושא): 'בא לי לבשל פסטה' ← סדרות בישול אמיתיות "
        "(The Bear, The Great British Baking Show, Chef's Table); סדרות שמזכירות אוכל אגב "
        "(כמו Queer Eye) מסוננות החוצה.",
        "מקרה כשל (בדוק ביחידת בדיקה): בקשה לסדרות דומות לזרע שאינו קיים בקטלוג ← המנוע "
        "מחזיר רשימה ריקה והשרת מחזיר הודעה חיננית ('לא נמצאו המלצות'), בלי לקרוס ובלי "
        "להמציא כותר (test_seeds.py).",
        "ג'יבריש (הקשה אקראית) ← בקשה מנומסת לנסח מחדש.",
    ])

    # 8
    heading(doc, "8. שיקולים אתיים", level=1, bookmark="sec8")
    bullets(doc, [
        "הטיית שפה: 62% מהקטלוג באנגלית ורק כ-10 כותרים בעברית; הצבנו ציפיות בהתאם והמנוע "
        "ניטרלי-שפה (Jaccard) מקבל משקל גבוה יותר בשאילתות בעברית.",
        "הטיית פופולריות ודירוג: דירוגים גבוהים מבוססי מעט הצבעות מנופחים; טיפלנו בכך עם דירוג "
        "IMDb משוקלל-הצבעות, כך שכותרים לא-מאומתים אינם מנצחים שלא בצדק.",
        "שקיפות ועיגון: ההמלצות מגיעות אך ורק מהקטלוג (catalog-first) - מניעת הזיות והטעיה.",
        "פרטיות: השיחה אינה נשמרת בשרת; אין איסוף מידע אישי מזהה.",
        "מגבלת איכות בעברית של מודל קטן - מנוטרת ומטופלת בהנחיות ובמנגנוני גיבוי.",
    ])

    # 9
    heading(doc, "9. מסקנות, יישום ומגבלות", level=1, bookmark="sec9")
    para(doc, "מסקנות:", bold=True)
    bullets(doc, [
        "דירוג משוקלל שקוף עדיף על אשכול שקרס עבור התאמה אישית אמיתית וניתנת להסבר.",
        "המתאם הנמוך בין מדדי הדמיון מצדיק גישה היברידית עבור 'דומה ל-X'.",
    ])
    para(doc, "יישום בעולם האמיתי: מנוע המלצות לפלטפורמות סטרימינג, ספריות תוכן, וצ'אטבוטים של VOD.")
    para(doc, "הנחות ומגבלות:", bold=True)
    bullets(doc, [
        "הקטלוג סטטי (snapshot) ואינו מתעדכן בזמן אמת.",
        "כיסוי עברי דל בקטלוג מגביל המלצות לתוכן ישראלי.",
        "אין חישוב שיבוצים בזמן ריצה (אילוץ משאבים), ולכן ה'גישור' מבוסס מילות מפתח.",
        "כ-5% מהתקצירים עדיין חסרים לאחר ההשלמה (סדרות חדשות מאוד שאינן בתמונת המצב של "
        "tvs.csv); עבורן השיבוץ מתבסס על הכותרת.",
        "תלות ב-Groq free tier (מגבלות קצב) ובשירות TMDB להעשרה (פוסטרים, טריילרים).",
    ])

    # 10
    heading(doc, "10. הערכת השפעה", level=1, bookmark="sec10")
    para(doc, "יתרונות פוטנציאליים:", bold=True)
    bullets(doc, [
        "התאמה אישית מהירה (3-4 שאלות ← 1-3 המלצות) המקצרת את זמן ה'מה לראות'.",
        "שקיפות ועיגון בקטלוג שמפחיתים מידע שגוי והזיות.",
        "נגישות דו-לשונית (עברית/אנגלית) לקהל הישראלי.",
    ])
    para(doc, "אתגרים אפשריים:", bold=True)
    bullets(doc, [
        "קנה מידה ורעננות הקטלוג בסביבת ייצור.",
        "שיפור איכות העברית והרחבת הכיסוי הלשוני.",
        "ניהול עלויות ומגבלות קצב של מודל השפה בסקיילה.",
    ])

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
